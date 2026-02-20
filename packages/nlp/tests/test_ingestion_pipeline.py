"""
Tests for the complete ingestion pipeline.
"""

from __future__ import annotations

import io
from pathlib import Path

import httpx
import pytest

from forge_nlp.extractors.rule_based import EntityAnnotation
from forge_nlp.pipeline.contract_metadata_mapper import ContractMetadata, map_entities_to_metadata
from forge_nlp.pipeline.ingestion_pipeline import (
    InMemoryDbClient,
    IngestionPipeline,
    LocalFileS3Client,
    extract_text_from_docx,
)
from forge_nlp.pipeline.quality_checker import (
    IssueSeverity,
    QualityReport,
    check_quality,
)

_FIXTURES = Path(__file__).parent / "fixtures"


# ─── Fixtures ─────────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def sample_docx_bytes() -> bytes:
    return (_FIXTURES / "sample_contract.docx").read_bytes()


@pytest.fixture(scope="module")
def sample_text(sample_docx_bytes: bytes) -> str:
    return extract_text_from_docx(sample_docx_bytes)


@pytest.fixture()
def s3_client() -> LocalFileS3Client:
    return LocalFileS3Client(base_dir=_FIXTURES)


@pytest.fixture()
def db_client() -> InMemoryDbClient:
    return InMemoryDbClient()


@pytest.fixture()
def pipeline(s3_client, db_client) -> IngestionPipeline:
    return IngestionPipeline(
        s3_client=s3_client,
        db_client=db_client,
        s3_bucket="test",
        use_ner=False,  # rule-based only for speed
        model_version="test-v0.1",
    )


# ═══════════════════════════════════════════════════════════════════════
# Text extraction tests
# ═══════════════════════════════════════════════════════════════════════


class TestTextExtraction:
    def test_docx_extracts_text(self, sample_text: str):
        """Text extraction from .docx should preserve content."""
        assert len(sample_text) > 500
        assert "FA8726-24-C-0042" in sample_text
        assert "52.202-1" in sample_text

    def test_docx_preserves_sections(self, sample_text: str):
        """Section headers should be present in the extracted text."""
        assert "SECTION A" in sample_text
        assert "SECTION B" in sample_text
        assert "SECTION I" in sample_text


# ═══════════════════════════════════════════════════════════════════════
# Full pipeline tests
# ═══════════════════════════════════════════════════════════════════════


class TestFullPipeline:
    def test_pipeline_with_sample_docx(self, pipeline: IngestionPipeline, db_client: InMemoryDbClient):
        """Full pipeline should process the sample .docx end-to-end."""
        result = pipeline.ingest(s3_key="sample_contract.docx", document_type="docx")

        assert result.contract_id is not None
        assert result.s3_key == "sample_contract.docx"
        assert result.document_type == "docx"
        assert result.text_length > 500
        assert result.chunk_count > 0
        assert result.entity_count > 0
        assert result.chunks_stored > 0
        assert result.duration_ms >= 0

    def test_contract_record_created(self, pipeline: IngestionPipeline, db_client: InMemoryDbClient):
        """A contract record should be created in the DB."""
        result = pipeline.ingest(s3_key="sample_contract.docx", document_type="docx")

        assert len(db_client.contracts) == 1
        contract = db_client.contracts[result.contract_id]
        assert contract["contract_number"] == "FA8726-24-C-0042"
        assert contract["s3_document_key"] == "sample_contract.docx"

    def test_chunks_stored_with_embeddings(self, pipeline: IngestionPipeline, db_client: InMemoryDbClient):
        """Chunks should be stored with 768-dim embeddings."""
        result = pipeline.ingest(s3_key="sample_contract.docx", document_type="docx")

        assert len(db_client.chunks) == result.chunks_stored
        for chunk_data in db_client.chunks.values():
            assert chunk_data["contract_id"] == result.contract_id
            assert chunk_data["document_s3_key"] == "sample_contract.docx"
            assert len(chunk_data["embedding"]) == 768
            assert chunk_data["chunk_text"] != ""

    def test_entity_annotations_linked_to_chunks(
        self, pipeline: IngestionPipeline, db_client: InMemoryDbClient,
    ):
        """Entity annotations should be stored and linked to chunk IDs."""
        result = pipeline.ingest(s3_key="sample_contract.docx", document_type="docx")

        assert result.annotations_stored > 0
        assert len(db_client.annotations) == result.annotations_stored

        # All annotations should reference valid chunk IDs
        valid_chunk_ids = set(db_client.chunks.keys())
        for ann in db_client.annotations:
            assert ann["chunk_id"] in valid_chunk_ids
            assert ann["model_version"] == "test-v0.1"

    def test_audit_log_entry_created(self, pipeline: IngestionPipeline, db_client: InMemoryDbClient):
        """Audit log entries should be created for the ingestion run."""
        pipeline.ingest(s3_key="sample_contract.docx", document_type="docx")

        # Should have at least 2 entries: RUNNING + SUCCESS/NEEDS_REVIEW
        assert len(db_client.audit_logs) >= 2
        types = [l["agent_type"] for l in db_client.audit_logs]
        assert all(t == "ingestion_pipeline" for t in types)

        statuses = [l["status"] for l in db_client.audit_logs]
        assert "RUNNING" in statuses
        assert any(s in ("SUCCESS", "NEEDS_REVIEW") for s in statuses)

    def test_pipeline_handles_no_sections_gracefully(
        self, db_client: InMemoryDbClient,
    ):
        """A document with no UCF sections should still be processed."""
        # Create a minimal docx with no section headers
        import docx
        doc = docx.Document()
        doc.add_paragraph("This is a plain contract with no section headers.")
        doc.add_paragraph("Contract number: FA8750-25-C-0100.")
        doc.add_paragraph("The total value is $1,000,000.")

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # Write to a temp location for S3 mock
        import tempfile
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "plain.docx"
            path.write_bytes(docx_bytes)

            s3 = LocalFileS3Client(base_dir=tmpdir)
            pipe = IngestionPipeline(
                s3_client=s3,
                db_client=db_client,
                use_ner=False,
            )
            result = pipe.ingest(s3_key="plain.docx", document_type="docx")

        assert result.chunk_count >= 1
        assert result.entity_count >= 1


# ═══════════════════════════════════════════════════════════════════════
# Metadata mapper tests
# ═══════════════════════════════════════════════════════════════════════


class TestMetadataMapper:
    def test_maps_contract_number(self):
        """CONTRACT_NUMBER entity should map to contract_number field."""
        entities = [
            EntityAnnotation("CONTRACT_NUMBER", "FA8726-24-C-0042", 0, 16, 1.0),
        ]
        meta = map_entities_to_metadata("Contract: FA8726-24-C-0042", entities)
        assert meta.contract_number == "FA8726-24-C-0042"

    def test_maps_naics_and_psc(self):
        """NAICS and PSC entities should map correctly."""
        entities = [
            EntityAnnotation("NAICS_CODE", "541330", 0, 6, 1.0),
            EntityAnnotation("PSC_CODE", "R425", 10, 14, 1.0),
        ]
        meta = map_entities_to_metadata("541330 code R425 code", entities)
        assert meta.naics_code == "541330"
        assert meta.psc_code == "R425"

    def test_maps_ceiling_value_with_context(self):
        """DOLLAR_AMOUNT near 'ceiling' context should map to ceiling_value."""
        text = "The total ceiling value of this contract is $12,500,000."
        entities = [
            EntityAnnotation("DOLLAR_AMOUNT", "$12,500,000", 43, 54, 1.0, {"normalized": "12500000"}),
        ]
        meta = map_entities_to_metadata(text, entities)
        assert meta.ceiling_value == "12500000"

    def test_maps_funded_value_with_context(self):
        """DOLLAR_AMOUNT near 'funded' context should map to funded_value."""
        text = "The contract is currently funded in the amount of $5,000,000."
        entities = [
            EntityAnnotation("DOLLAR_AMOUNT", "$5,000,000", 49, 59, 1.0, {"normalized": "5000000"}),
        ]
        meta = map_entities_to_metadata(text, entities)
        assert meta.funded_value == "5000000"

    def test_maps_pop_dates_from_pop_range(self):
        """POP_RANGE entities should set pop_start and pop_end."""
        text = "Period of performance: 01 January 2024 through 31 December 2025."
        entities = [
            EntityAnnotation(
                "POP_RANGE",
                "01 January 2024 through 31 December 2025",
                22, 62, 1.0,
                {"start_date": "2024-01-01", "end_date": "2025-12-31"},
            ),
        ]
        meta = map_entities_to_metadata(text, entities)
        assert meta.pop_start == "2024-01-01"
        assert meta.pop_end == "2025-12-31"

    def test_maps_security_level(self):
        """SECURITY_LEVEL entity should map correctly."""
        entities = [
            EntityAnnotation("SECURITY_LEVEL", "SECRET", 0, 6, 1.0),
        ]
        meta = map_entities_to_metadata("Classification: SECRET", entities)
        assert meta.security_level == "SECRET"

    def test_maps_far_and_dfars_clauses(self):
        """FAR and DFARS clause entities should be collected."""
        entities = [
            EntityAnnotation("FAR_CLAUSE", "52.202-1", 0, 8, 1.0),
            EntityAnnotation("FAR_CLAUSE", "52.212-4", 10, 18, 1.0),
            EntityAnnotation("DFARS_CLAUSE", "252.204-7012", 20, 32, 1.0),
        ]
        meta = map_entities_to_metadata("52.202-1 x 52.212-4 x 252.204-7012", entities)
        assert "52.202-1" in meta.far_clauses
        assert "52.212-4" in meta.far_clauses
        assert "252.204-7012" in meta.dfars_clauses

    def test_full_sample_docx_metadata(self, sample_text: str):
        """Metadata from the sample contract should include key fields."""
        from forge_nlp.extractors.rule_based import extract_all_entities

        entities = extract_all_entities(sample_text)
        meta = map_entities_to_metadata(sample_text, entities)

        assert meta.contract_number == "FA8726-24-C-0042"
        assert meta.naics_code == "541330"
        assert meta.psc_code == "R425"
        assert meta.cage_code == "1ABC2"
        assert len(meta.far_clauses) >= 2
        assert len(meta.dfars_clauses) >= 1


# ═══════════════════════════════════════════════════════════════════════
# Quality checker tests
# ═══════════════════════════════════════════════════════════════════════


class TestQualityChecker:
    def test_flags_missing_contract_number(self):
        """Missing contract_number should produce an ERROR issue."""
        meta = ContractMetadata()  # all fields None
        entities: list[EntityAnnotation] = []
        report = check_quality(meta, entities)

        codes = [i.code for i in report.issues]
        assert "MISSING_CONTRACT_NUMBER" in codes
        assert report.needs_human_review is True

    def test_flags_missing_ceiling_value(self):
        """Missing ceiling_value should produce a WARNING."""
        meta = ContractMetadata(contract_number="FA8726-24-C-0042")
        entities = [EntityAnnotation("CONTRACT_NUMBER", "FA8726-24-C-0042", 0, 16, 1.0)]
        report = check_quality(meta, entities)

        codes = [i.code for i in report.issues]
        assert "MISSING_CEILING_VALUE" in codes

    def test_flags_conflicting_contract_numbers(self):
        """Two different contract numbers should produce an ERROR."""
        meta = ContractMetadata(contract_number="FA8726-24-C-0042")
        entities = [
            EntityAnnotation("CONTRACT_NUMBER", "FA8726-24-C-0042", 0, 16, 1.0),
            EntityAnnotation("CONTRACT_NUMBER", "W911NF-23-D-0017", 50, 66, 1.0),
        ]
        report = check_quality(meta, entities)

        codes = [i.code for i in report.issues]
        assert "CONFLICTING_CONTRACT_NUMBERS" in codes
        assert report.needs_human_review is True

    def test_flags_low_confidence_entities(self):
        """Entities with confidence < 0.6 should be flagged."""
        meta = ContractMetadata(contract_number="X")
        entities = [
            EntityAnnotation("CONTRACTING_OFFICER", "John Smith", 0, 10, 0.3, {"source": "ner"}),
            EntityAnnotation("SCOPE_DESCRIPTION", "some work", 20, 29, 0.4, {"source": "ner"}),
        ]
        report = check_quality(meta, entities)

        codes = [i.code for i in report.issues]
        assert "LOW_CONFIDENCE_ENTITIES" in codes

    def test_flags_empty_chunks(self):
        """Many chunks with no entities should trigger a warning."""
        meta = ContractMetadata(contract_number="X")
        entities = [EntityAnnotation("CONTRACT_NUMBER", "X", 0, 1, 1.0)]
        report = check_quality(
            meta, entities,
            chunk_count=10,
            chunk_entity_counts=[0, 0, 0, 0, 0, 0, 1, 0, 0, 0],  # 9 of 10 empty
        )

        codes = [i.code for i in report.issues]
        assert "MANY_EMPTY_CHUNKS" in codes
        assert report.needs_human_review is True

    def test_no_issues_for_complete_metadata(self):
        """Complete metadata should produce no ERROR issues."""
        meta = ContractMetadata(
            contract_number="FA8726-24-C-0042",
            ceiling_value="12500000",
            funded_value="5000000",
            pop_start="2024-01-01",
            pop_end="2025-12-31",
            naics_code="541330",
        )
        entities = [
            EntityAnnotation("CONTRACT_NUMBER", "FA8726-24-C-0042", 0, 16, 1.0),
            EntityAnnotation("DOLLAR_AMOUNT", "$12,500,000", 20, 31, 1.0),
        ]
        report = check_quality(meta, entities)

        errors = [i for i in report.issues if i.severity == IssueSeverity.ERROR]
        assert len(errors) == 0
        assert report.needs_human_review is False


# ═══════════════════════════════════════════════════════════════════════
# API endpoint test
# ═══════════════════════════════════════════════════════════════════════


class TestAPIEndpoint:
    @pytest.fixture(scope="class")
    async def client(self, tmp_path_factory):
        """Create async client with the pipeline endpoint configured."""
        import os
        import api as api_module

        # Point S3_LOCAL_DIR to fixtures
        os.environ["S3_LOCAL_DIR"] = str(_FIXTURES)

        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=api_module.app),
            base_url="http://test",
        ) as c:
            yield c

        os.environ.pop("S3_LOCAL_DIR", None)

    @pytest.mark.asyncio
    async def test_ingest_endpoint(self, client: httpx.AsyncClient):
        """POST /pipeline/ingest should return structured results."""
        resp = await client.post("/pipeline/ingest", json={
            "s3_key": "sample_contract.docx",
            "document_type": "docx",
        })
        assert resp.status_code == 200
        data = resp.json()

        assert "result" in data
        assert "quality" in data
        assert data["result"]["chunk_count"] > 0
        assert data["result"]["entity_count"] > 0
        assert data["result"]["metadata"]["contract_number"] == "FA8726-24-C-0042"
